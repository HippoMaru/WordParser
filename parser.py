import docx 
import lxml.etree as ET
import copy


class Node:
    def __init__(self, id, tag, name, parent=None, value=None):
        self.id = id
        self.tag = tag
        self.name = name
        self.parent = parent
        self.value = value
        self.children = []


def parse_content_description_to_tree(doc):

    tocs = styles['toc 1'] + styles['toc 2'] + styles['toc 3']


    root = Node(0, 'main', 'ROOT')
    first_level_id = 1

    in_content_description = False
    for i in range(len(doc.paragraphs)):
        para = doc.paragraphs[i]
        print(para.style.name)
        if in_content_description:
            s_name = para.style.name

            if s_name not in tocs:
                return root
            if s_name in styles['toc 1']:
                node = Node(first_level_id, 'sec', ' '.join(para.text.split()[:-1]), root)
                root.children.append(node)
                first_level_id += 1
            else:
                id = para.text.split()[0].split('.')
                parent = root
                for i in range(len(id) - 1):
                    for ch in parent.children:
                        if ch.id == id[i]:
                            parent = ch
                            break
                node = Node(id[-1], 'subsec', ' '.join(para.text.split()[:-1]), parent)
                parent.children.append(node)

        if para.style.name in styles['content_description']:
            in_content_description = True


def parse_content_description_to_etree(doc):

    tocs = styles['toc 1'] + styles['toc 2'] + styles['toc 3']


    root = ET.Element('content_description')
    last_first_level = None
    last_second_level = None

    in_content_description = False
    for i in range(len(doc.paragraphs)):
        para = doc.paragraphs[i]
        print(para.style.name)
        if in_content_description:
            s_name = para.style.name

            if s_name not in tocs:
                return root

            if s_name in styles['toc 1']:
                node = ET.SubElement(root, 'sec')
                name = ' '.join(para.text.split()[:-1])
                if name[0].isdigit():
                    name = ' '.join(name.split()[1:])
                node.set('name', name)

                last_first_level = node

            if s_name in styles['toc 2']:
                node = ET.SubElement(last_first_level, 'subsec1')
                name = ' '.join(para.text.split()[:-1])
                if name[0].isdigit():
                    name = ' '.join(name.split()[1:])
                node.set('name', name)

                last_second_level = node

            if s_name in styles['toc 3']:
                node = ET.SubElement(last_second_level, 'subsec2')
                name = ' '.join(para.text.split()[:-1])
                if name[0].isdigit():
                    name = ' '.join(name.split()[1:])
                node.set('name', name)

        if para.style.name in styles['content_description']:
            in_content_description = True
    return root


def parse_em_all(doc):
    root = parse_content_description_to_etree(doc)
    prev = root

    tocs = styles['toc 1'] + styles['toc 2'] + styles['toc 3']

    levels = {
        'sec': 1,
        'subsec1': 2,
        'subsec2': 3,
    }

    doc_index = 0

    for elem in root.iter():
        if elem == root:
            continue
        if prev == root or levels[prev.tag] < levels[elem.tag]:
            prev = elem
            continue

        print(elem.get('name'))
        doc_index = 0

        while doc_index < len(doc.paragraphs):
            if prev.get('name') in doc.paragraphs[doc_index].text and doc.paragraphs[doc_index].style.name not in tocs:
                in_iter = False
                s_name = doc.paragraphs[doc_index].style.name
                iter_start = prev
                iter1_start = prev
                numpstart = prev
                doc_index += 1
                for i in range(doc_index, len(doc.paragraphs)):
                    para = doc.paragraphs[i]
                    print(para.style.name)
                    if para.style.name == s_name:
                        doc_index = i
                        break

                    if para.style.name in styles['nump 2']:
                        node = ET.SubElement(prev, 'levelledPara')
                        node.text = para.text
                        numpstart = node
                        continue
                    elif para.style.name in styles['nump 3']:
                        node = ET.SubElement(prev, 'levelledPara')
                        node.text = para.text
                        numpstart = node
                        continue
                    else:
                        numpstart = prev

                    if para.style.name in styles['iter 1']:
                        if not in_iter:
                            in_iter = True
                            iter_start = ET.SubElement(prev, 'sequentialList')
                        node = ET.SubElement(iter_start, 'listItem')
                        node.text = para.text
                        iter1_start = node
                        continue

                    if para.style.name in styles['iter 2']:
                        node = ET.SubElement(iter1_start, 'listItem')
                        node.text = para.text
                        continue

                    if para.style.name in styles['iter alt']:
                        if not in_iter:
                            in_iter = True
                            iter_start = ET.SubElement(prev, 'sequentialList')
                        if iter_start.tag in 'iter alt':
                            node = ET.SubElement(iter_start, 'listItem')
                        else:
                            node = ET.SubElement(iter1_start, 'listItem')
                        node.text = para.text
                        continue

                    in_iter = False

                    if para.style.name in styles['normal']:
                        node = ET.SubElement(numpstart, 'para')
                        node.text = para.text
                        continue

                    if styles['extra'][0] in para.style.name:
                        node = ET.SubElement(prev, 'extra')
                        node.text = para.text
                        numpstart = node

                    node = ET.SubElement(prev, para.style.name[:3])
                    node.text = para.text
                break
            doc_index += 1
        prev = elem

    return root

styles = {
    'content_description': ['Название-caps'],  # Содержание
    'toc 1': ['toc 1'],
    'toc 2': ['toc 2'],
    'toc 3': ['toc 3'],  # Перечисления в содержании (индексация длины 1, 2 и 3 соответственно)
    'subsec v': ['Нумерованный заголовок 3'],
    'iter 1': ["Перечисление-1"],
    'iter 2': ["Перечисление 2 уровень"],
    'iter alt': ["Перечень"],
    'nump 2': ["Нумерованный абзац 2"],
    'nump 3': ["Нумерованный абзац 3"],
    'normal': ["Normal"],
    'extra': ["Приложение"],
}

module_codees = {
    "DMC-VBMA-A-46-20-01-00A-018A-A_000_01_ru_RU.xml": "Введение",
    "DMC-VBMA-A-46-20-01-00A-020A-A_000_01_ru_RU.xml": "Назначение",
    "DMC-VBMA-A-46-20-01-00A-030A-A_000_01_ru_RU.xml": "Технические характеристики",
    "DMC-VBMA-A-46-20-01-00A-034A-A_000_01_ru_RU.xml": "Состав изделия",
    "DMC-VBMA-A-46-20-01-00A-041A-A_000_01_ru_RU.xml": "Устройство и работа",
    "DMC-VBMA-A-46-20-01-00A-044A-A_000_01_ru_RU.xml": "Описание и работа составных частей изделия",
    "DMC-VBMA-A-46-20-01-00A-122A-A_000_01_ru_RU.xml": "Указания по включению и опробованию работы изделия",
    "DMC-VBMA-A-46-20-01-00A-123A-A_000_01_ru_RU.xml": "Установка и настройка программного обеспечения",
    "DMC-VBMA-A-46-20-01-00A-410A-A_000_01_ru_RU.xml": "Перечень возможных неисправностей в процессе использования изделия и рекомендации по действиям при их возникновении",
}
# def run_parser(doc):
#     root = parse_em_all(doc)

#     et = ET.ElementTree(root)
#     et.write('output.xml', encoding="utf-8", pretty_print=True)

#     root = parse_content_description_to_etree(doc)
#     et = ET.ElementTree(root)
#     et.write('output2.xml', encoding="utf-8", pretty_print=True)
#     doc = docx.Document("input.docx")
    

######################################
##           TEST DRIVE             ##
######################################

def add_headers_to_content(content_xml_string, section_name: str):
    # Парсим content_xml_string в дерево элементов
    content_tree = ET.fromstring(content_xml_string)
    nsmap = {
        "rdf": "http://www.w3.org/1999/02/22-rdf-syntax-ns#",
        "ns2": "http://www.purl.org/dc/elements/1.1/",
        "xlink": "http://www.w3.org/1999/xlink",
        "xsi": "http://www.w3.org/2001/XMLSchema-instance"
    }

    # Create root element with nsmap
    dmodule = ET.Element("dmodule")
    # dmodule.set("xmlns:rdf", "http://www.w3.org/1999/02/22-rdf-syntax-ns#")
    dmodule.set("{{{}}}noNames".format(nsmap["rdf"]),
                "http://www.w3.org/1999/02/22-rdf-syntax-ns#")

    dmodule.set("{{{}}}noNames".format(nsmap["ns2"]),
                "http://www.purl.org/dc/elements/1.1/")

    dmodule.set("{{{}}}noNames".format(nsmap["xlink"]),
                "http://www.w3.org/1999/xlink")

    dmodule.set("{{{}}}noNames".format(nsmap["xsi"]),
                "http://www.w3.org/2001/XMLSchema-instance")

    dmodule.set("{{{}}}noNamespaceSchemaLocation".format(nsmap["xsi"]),
                "http://www.s1000d.org/S1000D_4-1/xml_schema_flat/descript.xsd")

    # Создаем элемент identAndStatusSection и его дочерние элементы
    q, w, e, r, t, y, u, i, o = [x for x in module_codees.keys() if module_codees[x] == section_name][0].split("-")
    a, s, d, f, g = o[:-4].split("_")
    ident_and_status_section = ET.SubElement(dmodule, "identAndStatusSection")
    dm_address = ET.SubElement(ident_and_status_section, "dmAddress")
    dm_ident = ET.SubElement(dm_address, "dmIdent")
    dm_code = ET.SubElement(dm_ident, "dmCode")
    # "DMC-VBMA-A-46-20-01-00A-041A-A_000_01_ru_RU.xml" a, s, d, f, g = o[:-4].split("_")
    dm_code.set("assyCode", y)
    dm_code.set("disassyCode", u[0:2])
    dm_code.set("disassyCodeVariant", e)
    dm_code.set("infoCode", i[0:3])
    dm_code.set("infoCodeVariant", e)
    dm_code.set("itemLocationCode", e)
    dm_code.set("modelIdentCode", w)
    dm_code.set("subSubSystemCode", t[1])
    dm_code.set("subSystemCode", t[0])
    dm_code.set("systemCode", r)
    dm_code.set("systemDiffCode", a)
    language = ET.SubElement(dm_ident, "language")
    language.set("countryIsoCode", g)
    language.set("languageIsoCode", f)
    issue_info = ET.SubElement(dm_ident, "issueInfo")
    issue_info.set("inWork", d)
    issue_info.set("issueNumber", s)
    dm_address_items = ET.SubElement(dm_address, "dmAddressItems")
    issue_date = ET.SubElement(dm_address_items, "issueDate")
    issue_date.set("day", "666")
    issue_date.set("month", "777")
    issue_date.set("year", "228")
    dm_title = ET.SubElement(dm_address_items, "dmTitle")
    tech_name = ET.SubElement(dm_title, "techName")
    tech_name.text = "Можно написать функцию, чтобы тут был другой текст"
    info_name = ET.SubElement(dm_title, "infoName")
    info_name.text = "Можно написать функцию, чтобы тут был другой текст"
    dm_status = ET.SubElement(ident_and_status_section, "dmStatus")
    security = ET.SubElement(dm_status, "security")
    responsible_partner_company = ET.SubElement(dm_status, "responsiblePartnerCompany")
    responsible_partner_company.set("enterpriseCode", "00000")
    originator = ET.SubElement(dm_status, "originator")
    originator.set("enterpriseCode", "00000")

    # Добавляем content в dmodule
    content = ET.SubElement(dmodule, "content")
    decription = ET.SubElement(content, "description")

    # Переносим дочерние элементы из content_tree в content

    for child in content_tree:
        decription.append(child)

    # Создаем дерево из dmodule и преобразуем его в XML-строку
    tree = ET.ElementTree(dmodule)
    return tree


# with open("/home/user/Desktop/test.xml", "rb") as file:
#     example = file.read()
# result_xml_string = add_headers_to_content(example, "Устройство и работа")
# print(result_xml_string)
# with open(r"/home/user/Desktop/result.xml", "w", encoding="utf-8") as f:
#   f.write(result_xml_string)

def extract_section_with_headers(doc, section_name: str):
    # Поиск subseca по сусекам
    iter = doc.getiterator(tag="subsec2")
    subsec_element = None
    for el in iter:
        if el.get('name') == section_name:
            subsec_element = el

    if subsec_element is None:
        raise ValueError("Subsec element not found for section: {}".format(section_name))

    # Конент
    content_element = ET.Element("content")
    content_element.append(subsec_element)

    # новое дерево с контентом
    content_tree = ET.ElementTree(content_element)
    

    # бл а
    content_xml_string = ET.tostring(content_tree.getroot(), encoding='unicode')
    result_et = add_headers_to_content(content_xml_string, section_name)

    return result_et


def run_parser(doc):
    doc = parse_em_all(doc)
    result_et = extract_section_with_headers(doc, "Технические характеристики")
    result_et.write('asshole.xml', encoding="utf-8", pretty_print=True)


doc = docx.Document("input.docx")
run_parser(doc)

# type: ignore